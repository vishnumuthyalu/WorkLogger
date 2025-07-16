import pandas as pd
from docx import Document
from io import BytesIO

def convert_df_to_csv(df):
    return df.to_csv(index=False).encode("utf-8")

def convert_df_to_docx(df, log_date):
    doc = Document()
    doc.add_heading("Daily Work Log", level=1)
    doc.add_paragraph(f"Date: {log_date.strftime('%A, %B %d, %Y')}")

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
