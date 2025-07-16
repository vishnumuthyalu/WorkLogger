import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from datetime import datetime, date
import smtplib
import ssl
from email.message import EmailMessage

st.set_page_config(page_title="Daily Work Log", layout="centered")
st.title("🗓️ Daily Work Log")

# ------------------------------------------------------------------
# Email Configuration (overridden by st.secrets if present)
# ------------------------------------------------------------------
_email_cfg = st.secrets.get("email", {}) if hasattr(st, "secrets") else {}

SMTP_SERVER = _email_cfg.get("server", "smtp.gmail.com")
SMTP_PORT = int(_email_cfg.get("port", 465))  # SSL default
SMTP_USER = _email_cfg.get("user", "your_email@example.com")
SMTP_PASSWORD = _email_cfg.get("password", "CHANGE_ME")  # Use st.secrets!
SENDER_NAME = _email_cfg.get("sender_name", "Daily Work Logger")
DEFAULT_TO = _email_cfg.get("default_to", "")
DEFAULT_CC = _email_cfg.get("default_cc", "")
DEFAULT_SUBJECT_TEMPLATE = _email_cfg.get("default_subject", "{date_str} Daily Work Log")

# ------------------------------------------------------------------
# DATE PICKER (no "Use Today" button here yet; user may add separately)
# ------------------------------------------------------------------
log_date = st.date_input("📆 Select the date to log for", value=date.today())
file_date_str = log_date.strftime('%A_%B_%d_%Y')  # Used for filenames

# Display current time
current_time = datetime.now()
st.markdown(f"**📅 Logging for:** `{file_date_str}`")
st.markdown(f"**⏰ Current Time:** `{current_time.strftime('%I:%M %p')}`")
st.divider()

# Work hours: 8 AM to 5 PM
hours = [f"{h}:00 {'AM' if h < 12 else 'PM'}" for h in range(8, 18)]

# Unique session key per date
log_key = f"structured_log_{log_date}"

# Initialize structured log for this date
if log_key not in st.session_state:
    st.session_state[log_key] = {
        hour: {
            "meeting": False,
            "meeting_info": "",
            "tasks": "",
            "general": ""
        } for hour in hours
    }

st.markdown("Fill out your log below for each hour:")

# Input form per hour
for hour in hours:
    with st.expander(f"🕒 {hour}"):
        st.session_state[log_key][hour]["meeting"] = st.checkbox(
            "Was there a meeting during this hour?",
            key=f"{log_key}_{hour}_meeting"
        )
        if st.session_state[log_key][hour]["meeting"]:
            st.session_state[log_key][hour]["meeting_info"] = st.text_area(
                "📝 Meeting Information", key=f"{log_key}_{hour}_meeting_info"
            )
        st.session_state[log_key][hour]["tasks"] = st.text_area(
            "✅ Tasks Worked On", key=f"{log_key}_{hour}_tasks"
        )
        st.session_state[log_key][hour]["general"] = st.text_area(
            "🗒️ General Information", key=f"{log_key}_{hour}_general"
        )

# Convert log to DataFrame
records = []
for hour in hours:
    entry = st.session_state[log_key][hour]
    records.append({
        "Time": hour,
        "Meeting": "Yes" if entry["meeting"] else "No",
        "Meeting Information": entry["meeting_info"] if entry["meeting"] else "",
        "Tasks": entry["tasks"],
        "General Information": entry["general"]
    })

log_df = pd.DataFrame(records)

# Export to CSV (bytes)
def convert_df_to_csv(df):
    return df.to_csv(index=False).encode("utf-8")

# Export to Word Document (BytesIO)
def convert_df_to_docx(df, log_date):
    doc = Document()
    doc.add_heading("Daily Work Log", level=1)
    doc.add_paragraph(f"Date: {log_date.strftime('%A, %B %d, %Y')}")
    
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Generate export bytes once so we can reuse for downloads *and* email
csv_bytes = convert_df_to_csv(log_df)
docx_io = convert_df_to_docx(log_df, log_date)
docx_bytes = docx_io.getvalue()

st.divider()

# ------------------------------------------------------------------
# DOWNLOAD BUTTONS
# ------------------------------------------------------------------
col1, col2 = st.columns(2)

with col1:
    st.download_button(
        "📤 Download as CSV",
        data=csv_bytes,
        file_name=f"{file_date_str}_daily_work_log.csv",
        mime="text/csv",
    )

with col2:
    st.download_button(
        "📄 Download as Word Document",
        data=docx_bytes,
        file_name=f"{file_date_str}_daily_work_log.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.success("✅ Your log is ready to download once you're done filling it out!", icon="✅")

st.divider()

# ------------------------------------------------------------------
# EMAIL SECTION
# ------------------------------------------------------------------
st.subheader("📧 Email This Work Log")

# Build default subject from template
default_subject = DEFAULT_SUBJECT_TEMPLATE.format(date_str=file_date_str)

# Email inputs
to_input = st.text_input("To", value=DEFAULT_TO, placeholder="recipient@example.com")
cc_input = st.text_input("CC (optional)", value=DEFAULT_CC, placeholder="cc1@example.com; cc2@example.com")
subject_input = st.text_input("Subject", value=default_subject)
from_display = st.text_input("From Name (display)", value=SENDER_NAME, help="Shown in recipient inbox display name.")
from_email_display = st.text_input("From Email (configured for SMTP login)", value=SMTP_USER, disabled=True)

# Email body preview (editable if you like)
default_body = f"Please find attached the work log for {log_date.strftime('%A, %B %d, %Y')}."
body_input = st.text_area("Body", value=default_body, height=100)

# Send button
if st.button("📧 Send Email with Attachments"):
    if not to_input.strip():
        st.error("Please provide at least one recipient email address.")
    elif SMTP_PASSWORD == "CHANGE_ME":
        st.error("SMTP password not configured. Set credentials in st.secrets before sending.")
    else:
        try:
            # Prepare message
            msg = EmailMessage()
            msg["Subject"] = subject_input
            msg["From"] = f"{from_display} <{SMTP_USER}>"
            msg["To"] = to_input

            # Support multiple CC separated by ; or ,
            cc_clean = [e.strip() for e in cc_input.replace(";", ",").split(",") if e.strip()]
            if cc_clean:
                msg["Cc"] = ", ".join(cc_clean)

            msg.set_content(body_input)

            # Attach CSV
            msg.add_attachment(
                csv_bytes,
                maintype="text",
                subtype="csv",
                filename=f"{file_date_str}_daily_work_log.csv",
            )

            # Attach DOCX
            msg.add_attachment(
                docx_bytes,
                maintype="application",
                subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{file_date_str}_daily_work_log.docx",
            )

            # SMTP send
            context = ssl.create_default_context()
            with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT, context=context) as server:
                server.login(SMTP_USER, SMTP_PASSWORD)
                all_recipients = [to_input] + cc_clean
                server.send_message(msg, from_addr=SMTP_USER, to_addrs=all_recipients)

            st.success("Email sent successfully! ✅")

        except Exception as e:
            st.error(f"Failed to send email: {e}")
